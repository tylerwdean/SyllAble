#requires python-docx. This can be installed with pip install python-docx

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches
from docx.shared import Pt
import json
import sys
import os

input_file_path = sys.argv[1]
output_file_path = sys.argv[2]

script_dir = os.path.dirname(os.path.abspath(__file__))

# Construct the path to FUS_Logo.png
logo_path = os.path.join(script_dir, "FUS_Logo.png")

def add_paragraph(document, paragraph):
    if paragraph['style'] == 'bullet':
        add_bullet_paragraph(document, paragraph)
    if paragraph['style'] == 'table':
        add_table_paragraph(document, paragraph)
    elif paragraph['style'] == 'text':
        add_text_paragraph(document, paragraph)

def add_text_paragraph(document, paragraph):
    #Creates the title of the new paragraph
    para = document.add_paragraph()
    run = para.add_run(paragraph['title'])
    run.bold = True
    run.font.size = Pt(14)
    
    #Adds the body message
    run = para.add_run(f"\n{paragraph['content']}")
    run.font.size = Pt(14)


def add_table_paragraph(document, paragraph):
    
    #Add a title if present for the table
    if len(paragraph['title']) > 0:
        para = document.add_paragraph()
        run = para.add_run(paragraph['title'])
        run.bold = True
        run.font.size = Pt(14)

    #max number of col is 5
    rows = paragraph['rows']
    rowNum=len(rows)
    colNum=len(rows[0])

    table = document.add_table(rows=rowNum, cols=colNum)
    table.style = 'Table Grid'
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    #The algorithm will first get the average length of the text in each column
    averageItemInColLen = []  #This has the average len() of the strings in each col, as an int
    for i in range(colNum):
        sumRowLen = 0
        for j in range(rowNum):
            sumRowLen += len(rows[j][i])
        averageItemInColLen.append(int(sumRowLen/rowNum)+1)
    
    #finds the longest len() item and adds that to the array. Used for resizing the table at the end.
    longestLenInCol = []
    for i in range(colNum):
        indexOfLongest=0
        for j in range(rowNum):
            if len(rows[j][i]) > len(rows[indexOfLongest][i]):
                indexOfLongest = j
        longestLenInCol.append(len(rows[indexOfLongest][i]))

    #Calculate the ratio of each number to each other to get percentage of total space to take up. 
    total = 0
    ratio=[]
    for i in range(colNum):
        total += averageItemInColLen[i]
    for i in range(colNum):
        #array of floating point numbers representing the ratio
        ratio.append(averageItemInColLen[i]/total)
    print(ratio)
    
    #From there, we will get the sizes of each col, but we will also make sure they are a min size of 10 characters
    #The whole screen is 10080 units long, so multiply that by the ratio to get each one. 
    colWidths = []
    for i in range(colNum):
        colWidths.append(int(10080*ratio[i]))

    #if the col's width is less than the min width, take the remaining width from the largest col
    minColWidth = 1500
    for i in range(colNum):
        if (colWidths[i] < minColWidth):
            #calculate difference to get the small col to 10%
            diff = minColWidth-colWidths[i]
            colWidths[i] = minColWidth
            #Find the largest col and set the width to prev. width -diff
            maxIndex = colWidths.index(max(colWidths))
            #update the data array (used for easy calculations) and the real columns
            colWidths[maxIndex] -= diff

    #Shrink the columns to fit the text if the text is smaller than the whole screen
    #Instead of this being based on the average, like the ratio, this will shrink to the largest item
    #We will assume that each character will take .7 of a character, and we are using 14pt font size
    for i in range(colNum):
        #This works for something of small length, like 2 or 3
        textWidth = 0
        if (longestLenInCol[i]<5):
            textWidth = int(longestLenInCol[i]*Pt(14).inches*1440*1.15) #text-length * (fontsize * 1440 to convert to twips)
        else:
            textWidth = int(longestLenInCol[i]*Pt(14).inches*1440*.7) #text-length * (fontsize * 1440 to convert to twips)
        
        if (colWidths[i] > textWidth):
            colWidths[i] = textWidth

    #Now that the table size is correct, we can insert the data into the table
    for i in range(colNum):
        for j in range(rowNum):
            cell = table.cell(j, i)
            cell.text = rows[j][i]
            cell.width = Inches(colWidths[i]/1440)

            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(14)
                    run.font.name = 'Times New Roman'
    
    table.width = Inches(sum(colWidths)/1440)
    document.add_paragraph()    #For spacing purposes


def add_bullet_paragraph(document, paragraph):
    if (len(paragraph['content']) > 0):
        reqItems = document.add_paragraph()
        run = reqItems.add_run(paragraph['title'])
        run.font.size = Pt(14)
        run.bold = True
        reqItems.paragraph_format.space_after = Pt(0)
    
    for item in paragraph['content']:
        para = document.add_paragraph(f"{item}", style='List Bullet')
        for run in para.runs:
            run.font.size = Pt(14)


f = open(input_file_path, encoding="utf8")
data = json.load(f)


document = Document()
section = document.sections[0]
section.left_margin = Inches(.75)    # Set left margin to 1 inch
section.right_margin = Inches(.75)   # Set right margin to 1 inch
section.top_margin = Inches(.75)     # Set top margin to 1 inch
section.bottom_margin = Inches(.75)  # Set bottom margin to 1 inch


image = document.add_paragraph()

run = image.add_run()
run.add_picture(logo_path, width=Inches(4))

image.alignment = WD_ALIGN_PARAGRAPH.CENTER

#Adds the course code to the top of the page
docTitle = document.add_paragraph()
run = docTitle.add_run("COURSE INFORMATION SHEET")
run.bold = True
run.font.size = Pt(20)
docTitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

courseTitle = document.add_paragraph()
run = courseTitle.add_run(data['courseCode'])
run.font.size = Pt(16)
courseTitle.alignment = WD_ALIGN_PARAGRAPH.CENTER



table = document.add_table(rows=1, cols=2)

# Set the width of the columns 
#table.columns[0].width = 3000000  # Width of the left column
#table.columns[1].width = 3000000  # Width of the right column

# Fill the left column with professor's details
left_cell = table.cell(0, 0)
left_cell.text = f"Professor: {data['professorName']}\nOffice: {data['office']}\nOffice Hours: {data['officeHours']}\nPhone: {data['phone']}\nEmail: {data['email']}"

# Fill the right column with aligned items
right_cell = table.cell(0, 1)
right_cell.text = f"Semester: {data['semester']}\nClassroom: {data['classroom']}\nClass Meeting Days: {data['classDays']}\nClass Meeting Times: {data['classTimes']}"

# Adjust formatting for the items in the table
for para in right_cell.paragraphs:
    para.alignment = 2
    for run in para.runs:
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'

for para in left_cell.paragraphs:
    for run in para.runs:
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'

#The next thing it will write is the Course Description, using the same paragraph style as all other paragraphs

courseDescriptionParagraph = {
    "style": "text",
    "title": "Course Description",
    "content": data['courseDescription']
}
add_paragraph(document, courseDescriptionParagraph)

#Bulletted List of required items

#Now it will start generating 'normal' paragraphs
for paragraph in data['paragraphs']:
    add_paragraph(document, paragraph)
    

#Changes all the font to Times New Roman
for para in document.paragraphs:
    for run in para.runs:
        run.font.name = 'Times New Roman'

document.save(output_file_path)